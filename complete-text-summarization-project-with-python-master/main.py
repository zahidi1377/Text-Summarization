import nltk
import re
import string
import tkinter as tk
from tkinter.scrolledtext import ScrolledText
from bs4 import BeautifulSoup
from urllib.request import urlopen
import tkinter.filedialog as filedialog
from tkinter import messagebox, ttk
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from tkinter import PhotoImage
from PIL import Image, ImageTk

# start of algorithm
def preprocess(text):
    text = re.sub(r'\s+', ' ', text)
    stopwords = set(nltk.corpus.stopwords.words('english'))
    formatted_text = text.lower()
    tokens = nltk.word_tokenize(formatted_text)
    tokens = [word for word in tokens if word not in stopwords and word not in string.punctuation]
    formatted_text = ' '.join(tokens)
    return formatted_text
def calculate_sentence_scores(sentences, important_words, distance):
    scores = []
    for sentence in sentences:
        word_tokens = nltk.word_tokenize(sentence)
        word_indices = [i for i, word in enumerate(word_tokens) if word in important_words]
        if not word_indices:
            continue
        groups = []
        group = [word_indices[0]]
        for i in range(1, len(word_indices)):
            if word_indices[i] - word_indices[i - 1] < distance:
                group.append(word_indices[i])
            else:
                groups.append(group[:])
                group = [word_indices[i]]
        groups.append(group)  # Add the last group
        max_group_score = 0
        for group in groups:
            important_words_in_group = len(group)
            total_words_in_group = group[-1] - group[0] + 1
            score = (important_words_in_group ** 2) / total_words_in_group
            if score > max_group_score:
                max_group_score = score
        scores.append(max_group_score)
    return scores
def summarizer(text, percentage):
    original_sentences = nltk.sent_tokenize(text)
    formatted_text = preprocess(text)
    words = nltk.word_tokenize(formatted_text)
    frequency = nltk.FreqDist(words)
    top_n_words = [word for word, _ in frequency.most_common(5)]  # Top 5 words by frequency
    num_sentences = int(len(original_sentences) * percentage / 100)
    if num_sentences==0:
        percentage=int(percentage)
        messagebox.showerror("Error", "The text is too short. I can't summarize\n in given the percentage: " + str(percentage))
    scores = calculate_sentence_scores(original_sentences, top_n_words, distance=1)
    # Sort sentences by their original order and get the top 'num_sentences'
    ranked_indices = sorted(range(len(scores)), key=lambda k: original_sentences[k])
    best_sentences = [original_sentences[i] for i in ranked_indices[:num_sentences]]
    return best_sentences
# end of algorithm
# create GUI window
window = tk.Tk()
window.title("Summarizer")
window.geometry('1000x700')
window.minsize(width=700, height=600)
custom_logo = tk.PhotoImage(file="AILOGO.png")
window.iconphoto(True, custom_logo)

# Create a style for ttk widgets (tabs)
# Add a padding style to the notebook tabs
padding_style = ttk.Style()
padding_style.configure('TNotebook.Tab', padding=[6, 3])
# create Tabs
tab_control = ttk.Notebook(window)
HomeTab1 = ttk.Frame(tab_control)
FileTab = ttk.Frame(tab_control)
URLTab = ttk.Frame(tab_control)
AboutUS_Tab = ttk.Frame(tab_control)
HelpTab = ttk.Frame(tab_control)
# add tabs to the tab control or top option
tab_control.add(HomeTab1, text=f'{"Home":^20s}')
tab_control.add(FileTab, text=f'{"File":^20s}')
tab_control.add(URLTab, text=f'{"URL":^20s}')
tab_control.add(AboutUS_Tab, text=f'{"About US":^20s}')
tab_control.add(HelpTab, text=f'{"Help":^20s}')
# configure row and column weights for responsive layout

for i in range(22):
    HomeTab1.grid_columnconfigure(i, weight=1)
    FileTab.grid_columnconfigure(i, weight=1)
    URLTab.grid_columnconfigure(i, weight=1)
HomeTab1.grid_rowconfigure(3, weight=1)
HomeTab1.grid_rowconfigure(4, weight=1)
FileTab.grid_rowconfigure(3, weight=1)
FileTab.grid_rowconfigure(4, weight=1)
URLTab.grid_rowconfigure(4, weight=1)
URLTab.grid_rowconfigure(5, weight=1)

# adding labels on Tabs uppering stexts
label1 = tk.Label(HomeTab1, text="Enter your texts to summarize.", padx=5, pady=5)
label1.grid(column=0, row=0)
label2 = tk.Label(FileTab, text="Open and select file to summarize.", padx=5, pady=5)
label2.grid(column=0, row=0)
label3 = tk.Label(URLTab, text="Write URL to get the text to summarize.", padx=5, pady=5)
label3.grid(column=0, row=0)

#buttons with added styles
button_style = {
    'bg': '#0055ff',
    'fg': '#fff',
    'cursor': 'hand2',
    'border': 7,
    'relief': 'flat',  # No border
    'font': ('Helvetica', 7, 'bold'),  # Bold font
    'width':'10',
}
# Function to apply hover effect on buttons
def on_enter(event):
    event.widget.config(bg='#001eff')
def on_leave(event):
    event.widget.config(bg='#0055ff')
# Start Home tab
def Home_clear_text():
    text = Home_Entry_Text.get('1.0', 'end-1c')
    if not text:
        pass
    else:
        Home_Entry_Text.delete('1.0', tk.END)
        messagebox.showinfo("Success", "clear successfully!")
def Home_clear_display_result():
    text = Home_display_Result_Text.get('1.0', 'end-1c')
    if not text:
        pass
    else:
        Home_display_Result_Text.delete('1.0', tk.END)
        HomeTab_percentage_Entry.delete(0,tk.END)
        messagebox.showinfo("Success", "clear successfully!")
def Home_save_summary():
    content = Home_display_Result_Text.get('1.0', 'end-1c')
    if not content:
        messagebox.showerror("Error", "No content to save.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[
        ("Text files", "*.txt"),
        ("PDF files", "*.pdf"),
        ("Word files", "*.docx"),
        ("All files", "*.*")
    ])
    if not file_path:
        return  # User canceled the save operation
    file_extension = file_path.split(".")[-1].lower()
    if file_extension == "txt":
        with open(file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(content)
        messagebox.showinfo("Success", "Text file saved successfully!")
    elif file_extension == "pdf":
        lines = content.split('\n')
        c = canvas.Canvas(file_path, pagesize=letter)
        y = 750  # Starting y-coordinate
        for line in lines:
            c.drawString(50, y, line)
            y -= 12  # Adjust this value for line spacing
        c.save()
        messagebox.showinfo("Success", "PDF file saved successfully!")
    elif file_extension == "docx":
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        messagebox.showinfo("Success", "Word document saved successfully!")
    else:
        messagebox.showerror("Error", "Unsupported file type. Please choose a valid file extension.")
def Home_get_summarize():
    percentage = HomeTab_percentage_Entry.get()  # Get the user-entered percentage
    if not percentage:
        messagebox.showerror("Error", "Please enter a percentage.")
        return
    try:
        percentage = float(percentage)
    except ValueError:
        messagebox.showerror("Error", "Invalid percentage format.")
        return
    if percentage <= 0 or percentage > 100:
        messagebox.showerror("Error", "Percentage should be between 0 and 100.")
        return
    text_to_summarize = Home_Entry_Text.get('1.0', 'end-1c')  # Get the text from the ScrolledText widget
    if not text_to_summarize:
        messagebox.showerror("Error", "Please enter text to summarize.")
        return
    best_sentences = summarizer(text_to_summarize, percentage)
    Home_display_Result_Text.delete('1.0', tk.END)  # Clear previous content
    summarized_text = "\n".join(best_sentences)
    Home_display_Result_Text.insert(tk.END, summarized_text)

# Label and Entry for percentage input HomeTab1
HomeTab_percentage_label = tk.Label(HomeTab1, text="Enter Percentage:", padx=5, pady=5,font=("Helvetica",10))
HomeTab_percentage_label.grid(row=1, column=2, padx=5, pady=5)
HomeTab_percentage_Entry = tk.Entry(HomeTab1)
HomeTab_percentage_Entry.grid(row=2, column=2, padx=5, pady=5)
# Buttons
Hbot1 = tk.Button(HomeTab1, text='Clear Text', command=Home_clear_text,**button_style)
Hbot1.grid(row=1, column=0, padx=5, pady=5)
Hbot1.bind("<Enter>", on_enter)
Hbot1.bind("<Leave>", on_leave)
Hbot2 = tk.Button(HomeTab1, text='Clear Result', command=Home_clear_display_result, **button_style)
Hbot2.grid(row=1, column=1, padx=5, pady=5)
Hbot2.bind("<Enter>", on_enter)
Hbot2.bind("<Leave>", on_leave)
Hbot3 = tk.Button(HomeTab1, text='Save', command=Home_save_summary, **button_style)
Hbot3.grid(row=2, column=0, padx=5, pady=5)
Hbot3.bind("<Enter>", on_enter)
Hbot3.bind("<Leave>", on_leave)
Hbot4 = tk.Button(HomeTab1, text='Summarize', command=Home_get_summarize, **button_style)
Hbot4.grid(row=2, column=1, padx=5, pady=5)
Hbot4.bind("<Enter>", on_enter)
Hbot4.bind("<Leave>", on_leave)
# Text area for type text for tab Home
Home_Entry_Text = ScrolledText(HomeTab1, height=22, width=250,font=("Helvetica",9))
Home_Entry_Text.grid(row=3, column=0, columnspan=18, pady=5, padx=5)
Home_Entry_Text.config(borderwidth=2, relief="groove")  # Added border and relief for improved design

# Text area to show summarization result
Home_display_Result_Text = ScrolledText(HomeTab1, height=22, width=250,font=("Helvetica",9))
Home_display_Result_Text.grid(row=4, column=0, columnspan=18, padx=5, pady=5)
Home_display_Result_Text.config(borderwidth=2, relief="groove")  # Added border and relief for improved design
# end HomeTab1 Home tab

# start of second tab file processing tab
# start of second tab file processing tab
def file_Open_file():
    file_path = filedialog.askopenfilename(
        filetypes=[('Text Files', '*.txt'), ('Word Files', '*.docx'), ('All Files', '*.*')])
    if file_path:
        file_extension = file_path.split('.')[-1].lower()
        if file_extension == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    read_text = file.read()
                    File_Entry_Text.delete('1.0', tk.END)  # Clear previous content
                    File_Entry_Text.insert(tk.END, read_text)
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as file:
                        read_text = file.read()
                        File_Entry_Text.delete('1.0', tk.END)  # Clear previous content
                        File_Entry_Text.insert(tk.END, read_text)
                except Exception as e:
                    messagebox.showerror("Error", f"An error occurred while opening the file: {e}")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while opening the file: {e}")
        elif file_extension == 'docx':
            try:
                doc = Document(file_path)
                read_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                File_Entry_Text.delete('1.0', tk.END)  # Clear previous content
                File_Entry_Text.insert(tk.END, read_text)
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while opening the Word file: {e}")
        else:
            messagebox.showerror("Error", "Unsupported file type. Please choose a valid file extension.")
def file_reset_tab():
    text=File_Entry_Text.get('1.0', 'end-1c')
    if not text:
        pass
    else:
        File_Entry_Text.delete('1.0', tk.END)
        File_display_Result_Text.delete('1.0', tk.END)
        FileTab_percentage_Entry.delete(0, tk.END)
        messagebox.showinfo("Success", "clear successfully!")
def file_save_summary():
    # Get the content from the ScrolledText widget
    content = File_display_Result_Text.get('1.0', 'end-1c')
    if not content:
        messagebox.showerror("Error", "No content to save.")
        return
    # Ask the user to choose the file type and location
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[
        ("Text files", "*.txt"),
        ("PDF files", "*.pdf"),
        ("Word files", "*.docx"),
        ("All files", "*.*")
    ])
    if not file_path:
        return  # User canceled the save operation
    # Determine the file extension
    file_extension = file_path.split(".")[-1].lower()
    if file_extension == "txt":
        # Save as a plain text file
        with open(file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(content)
        messagebox.showinfo("Success", "Text file saved successfully!")
    elif file_extension == "pdf":
        # Save as a PDF file
        lines = content.split('\n')
        c = canvas.Canvas(file_path, pagesize=letter)
        y = 750  # Starting y-coordinate
        for line in lines:
            c.drawString(50, y, line)
            y -= 12  # Adjust this value for line spacing
        c.save()
        messagebox.showinfo("Success", "PDF file saved successfully!")
    elif file_extension == "docx":
        # Save as a Word document
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        messagebox.showinfo("Success", "Word document saved successfully!")
    else:
        messagebox.showerror("Error", "Unsupported file type. Please choose a valid file extension.")
def file_get_summarize():
    percentage = FileTab_percentage_Entry.get()  # Get the user-entered percentage
    if not percentage:
        messagebox.showerror("Error", "Please enter a percentage.")
        return
    try:
        percentage = float(percentage)
    except ValueError:
        messagebox.showerror("Error", "Invalid percentage format.")
        return
    if percentage <= 0 or percentage > 100:
        messagebox.showerror("Error", "Percentage should be between 0 and 100.")
        return
    text_to_summarize = File_Entry_Text.get('1.0', 'end-1c')  # Get the text from the ScrolledText widget
    if not text_to_summarize:
        messagebox.showerror("Error", "Please enter text to summarize.")
        return
    best_sentences = summarizer(text_to_summarize, percentage)
    File_display_Result_Text.delete('1.0', tk.END)  # Clear previous content
    summarized_text = "\n".join(best_sentences)
    File_display_Result_Text.insert(tk.END, summarized_text)
# Label and Entry for percentage input FileTab file
Filetab_percentage_label = tk.Label(FileTab, text="Enter Percentage:", padx=5, pady=5,font=("Helvetica",10))
Filetab_percentage_label.grid(row=1, column=3, padx=5, pady=5)
FileTab_percentage_Entry = tk.Entry(FileTab)
FileTab_percentage_Entry.grid(row=2, column=3, padx=5, pady=5)
# Button
fBot1 = tk.Button(FileTab, text='Open File', command=file_Open_file,**button_style)
fBot1.grid(row=1, column=0, padx=5, pady=5)
fBot1.bind("<Enter>", on_enter)
fBot1.bind("<Leave>", on_leave)
fBot2 = tk.Button(FileTab, text='Reset', command=file_reset_tab,**button_style)
fBot2.grid(row=1, column=1, padx=5, pady=5)
fBot2.bind("<Enter>", on_enter)
fBot2.bind("<Leave>", on_leave)
fBot3 = tk.Button(FileTab, text='Save',command=file_save_summary, **button_style)
fBot3.grid(row=2, column=0, padx=5, pady=5)
fBot3.bind("<Enter>", on_enter)
fBot3.bind("<Leave>", on_leave)
fBot4 = tk.Button(FileTab, text='Summarize', command=file_get_summarize,**button_style)
fBot4.grid(row=2, column=1, padx=5, pady=5)
fBot4.bind("<Enter>", on_enter)
fBot4.bind("<Leave>", on_leave)
# text area for screen to show to put text to summarize Fil FileTab
File_Entry_Text = ScrolledText(FileTab, height=22, width=250,font=("Helvetica",9))
File_Entry_Text.grid(row=3, column=0, columnspan=18, padx=5, pady=5)
File_Entry_Text.config(borderwidth=2, relief="groove")
# text area for screen to show result Fil FileTab
File_display_Result_Text = ScrolledText(FileTab, height=22, width=250,font=("Helvetica",9))
File_display_Result_Text.grid(row=4, column=0, columnspan=18, padx=5, pady=5)
File_display_Result_Text.config(borderwidth=2, relief="groove")
# end FileTab file processing tab
# start if URLTab url tab
def CLear_url_text_display():
    text=URLTab_Entry_Text.get('1.0', 'end-1c')
    if not text:
        pass
    else:
        URLTab_Entry_Text.delete('1.0', tk.END)
        URLtab_display_Result_Text.delete('1.0', tk.END)
        url_entry.delete('0', tk.END)
        messagebox.showinfo("Success", "clear successfully!")
def URL_Get_text():
    raw_text = str(url_entry.get())
    if raw_text.strip():  # Check if the URL is not empty or whitespace
        try:
            page = urlopen(raw_text)
            soup = BeautifulSoup(page, 'lxml')
            fetched_text = ''.join(map(lambda p: p.text, soup.find_all('p')))
            URLTab_Entry_Text.delete('1.0', tk.END)  # Clear previous content
            URLTab_Entry_Text.insert(tk.END, fetched_text)  # Display fetched text in URLTab_Entry_Text
        except Exception as e:
            messagebox.showerror("Error",
                                "An error occurred while fetching text from url!")  # Show the error message in a new window
    else:
        messagebox.showerror("Error", "Please enter a valid URL!")
        # Show an error message in a new window
def URL_clear_display_result():
    text=URLtab_display_Result_Text.get('1.0', 'end-1c')
    if not text:
        pass
    else:
        URLtab_display_Result_Text.delete('1.0', tk.END)
        URLTab_percentage_Entry.delete(0, tk.END)
        url_entry.delete(0, tk.END)
        messagebox.showinfo("Success", "clear successfully!")
def URL_save_summary():
    # Get the content from the ScrolledText widget
    content = URLtab_display_Result_Text.get('1.0', 'end-1c')
    if not content:
        messagebox.showerror("Error", "No content to save.")
        return
    # Ask the user to choose the file type and location
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[
        ("Text files", "*.txt"),
        ("PDF files", "*.pdf"),
        ("Word files", "*.docx"),
        ("All files", "*.*")
    ])
    if not file_path:
        return  # User canceled the save operation
    # Determine the file extension
    file_extension = file_path.split(".")[-1].lower()
    if file_extension == "txt":
        # Save as a plain text file
        with open(file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(content)
        messagebox.showinfo("Success", "Text file saved successfully!")
    elif file_extension == "pdf":
        # Save as a PDF file
        lines = content.split('\n')
        c = canvas.Canvas(file_path, pagesize=letter)
        y = 750  # Starting y-coordinate
        for line in lines:
            c.drawString(50, y, line)
            y -= 12  # Adjust this value for line spacing
        c.save()
        messagebox.showinfo("Success", "PDF file saved successfully!")
    elif file_extension == "docx":
        # Save as a Word document
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        messagebox.showinfo("Success", "Word document saved successfully!")
    else:
        messagebox.showerror("Error", "Unsupported file type. Please choose a valid file extension.")
def URL_get_summarize():
    percentage = URLTab_percentage_Entry.get()  # Get the user-entered percentage
    if not percentage:
        messagebox.showerror("Error", "Please enter a percentage.")
        return
    try:
        percentage = float(percentage)
    except ValueError:
        messagebox.showerror("Error", "Invalid percentage format.")
        return
    if percentage <= 0 or percentage > 100:
        messagebox.showerror("Error", "Percentage should be between 0 and 100.")
        return
    text_to_summarize = URLTab_Entry_Text.get('1.0', 'end-1c')  # Get the text from the ScrolledText widget
    if not text_to_summarize:
        messagebox.showerror("Error", "Please enter text to summarize.")
        return
    best_sentences = summarizer(text_to_summarize, percentage)
    URLtab_display_Result_Text.delete('1.0', tk.END)  # Clear previous content
    summarized_text = "\n".join(best_sentences)
    URLtab_display_Result_Text.insert(tk.END, summarized_text)
# Label and Entry for percentage input FileTab file
URLtab_percentage_label = tk.Label(URLTab, text="Enter Percentage:", padx=5, pady=5,font=("Helvetica",10))
URLtab_percentage_label.grid(row=1, column=3, padx=5, pady=5)
URLTab_percentage_Entry = tk.Entry(URLTab)
URLTab_percentage_Entry.grid(row=2, column=3, padx=5, pady=5)
# getting url / third tab options url tab
li = tk.Label(URLTab, text="Enter Url to get text and summarize",font=("Helvetica",10))
li.grid(row=1, column=4)
# url text box
raw_entry = tk.StringVar()
url_entry = tk.Entry(URLTab, textvariable=raw_entry, width=50)
url_entry.grid(row=2, column=4)
# buttons
UBot0 = tk.Button(URLTab, text='Get Text', command=URL_Get_text, **button_style)
UBot0.grid(row=2, column=5, padx=5, pady=5)
UBot0.bind("<Enter>", on_enter)
UBot0.bind("<Leave>", on_leave)
UBot1 = tk.Button(URLTab, text="Clear Text", command=CLear_url_text_display, **button_style)
UBot1.grid(row=1, column=0, padx=5, pady=5)
UBot1.bind("<Enter>", on_enter)
UBot1.bind("<Leave>", on_leave)
UBot2 = tk.Button(URLTab, text='Clear Result', command=URL_clear_display_result,**button_style)
UBot2.grid(row=1, column=1, padx=5, pady=5)
UBot2.bind("<Enter>", on_enter)
UBot2.bind("<Leave>", on_leave)
UBot3 = tk.Button(URLTab, text="save", command=URL_save_summary, **button_style)
UBot3.grid(row=2, column=0, padx=5, pady=5)
UBot3.bind("<Enter>", on_enter)
UBot3.bind("<Leave>", on_leave)
UBot4 = tk.Button(URLTab, text='Summarize', command=URL_get_summarize, **button_style)
UBot4.grid(row=2, column=1, padx=5, pady=5)
UBot4.bind("<Enter>", on_enter)
UBot4.bind("<Leave>", on_leave)
# text area for screen to show to put text to summarize Fil FileTab
URLTab_Entry_Text = ScrolledText(URLTab, height=22, width=250,font=("Helvetica",9))
URLTab_Entry_Text.grid(row=3, column=0, columnspan=18, padx=5, pady=5)
URLTab_Entry_Text.config(borderwidth=2, relief="groove")  # Added border and relief for improved design

# text area for screen to show result Fil FileTab
URLtab_display_Result_Text = ScrolledText(URLTab, height=22, width=250,font=("Helvetica",9))
URLtab_display_Result_Text.grid(row=4, column=0, columnspan=18, padx=5, pady=5)
URLtab_display_Result_Text.config(borderwidth=2, relief="groove")
#end of url tab
#atart of about us tab
# Create About Us tab
logo_kateb = PhotoImage(file="logoKateb.png").subsample(1)
def fill_about_us_tab():
    about_us_text = """
    \t\tWelcome to the about us section!\n\n
    Project Name: Text Summarization
    Developed by: sarullah zahidi

    About the Project:
    The Text Summarization Tool is designed to summarize text from various sources, including user-provided text, files
    and web URLs. It uses natural language processing techniques to generate concise summaries, making it a valuable tool
    for extracting key information from large volumes of text.

    Contact Information:
    Email: sarullahzahidi@outlook.com
    GitHub: github.com/zahidi1377
    """
    # Create a label for the text
    about_us_label = tk.Label(AboutUS_Tab, text=about_us_text, padx=10, pady=10, justify="left", font=("Helvetica", 10))
    about_us_label.grid(row=2, column=0, padx=10, pady=10)
    # Create a label for the logo
    logokateb = tk.Label(AboutUS_Tab, image=logo_kateb)
    logokateb.grid(row=3, column=0, padx=5, pady=0, sticky="w")
# Call the function to fill the About Us tab
fill_about_us_tab()
# end tab about us
#start tab help HelpTab
def fill_Help_tab():
    help = """
\t\t\tWelcome to the help section!\n\n
This application allows you to summarize text from various sources. Here's how to use each tab:
1. Home Tab:
   - Enter the text you want to summarize in the text area.
   - Enter the desired summarization percentage.
   - Click the 'Summarize' button to generate a summary.
   - You can also save the summary to a file.

2. File Tab:
   - Click the 'Open File' button to select a text file for summarization.
   - You can also enter or paste text in the text area.
   - Follow the same steps as in the Home tab to summarize and save.

3. URL Tab:
   - Enter a valid URL in the URL text box.
   - Click the 'Get Text' button to fetch text from the URL.
   - Summarize and save the text as in the Home tab.

4. About US Tab:
   - Information about the application developer.

5. Help Tab:
   - You are here! This section provides information on how to use the application.

Note:
- Percentage should be between 0 and 100.
- Supported file formats for saving summaries are TXT, PDF, and DOCX.
    """

    help_label = tk.Label(HelpTab, text=help, padx=10, pady=10, justify="left",font=("Helvetica",10))
    help_label.grid(row=2, column=0, padx=10, pady=10)
# Call the function to fill the helps tab
fill_Help_tab()
# End help tab
# Pack the tab control to display it and start UI
tab_control.pack(expand=1, fill='both')
window.mainloop()
